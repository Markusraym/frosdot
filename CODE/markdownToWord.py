# Requires the permissively licensed python-docx, available on PyPI (https://pypi.org/project/python-docx/)

from docx import Document
import argparse
import re

parser = argparse.ArgumentParser()
parser.add_argument("fileName", help="Convert this file")
args = parser.parse_args()

document = Document()

# Title
title = re.sub(".md$", "", args.fileName)
document.add_heading(title)
myfile = open(args.fileName).read()
p = document.add_paragraph(myfile)

document.save("output.docx")
